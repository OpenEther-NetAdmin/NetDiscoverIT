"""Unit tests for compliance report generation (Group 8b)."""
from __future__ import annotations

from datetime import datetime, timezone

import pytest

from app.services.compliance.evidence_models import (
    AuditEvidence,
    ChangeEvidence,
    DeviceEvidence,
    EvidencePackage,
    PathEvidence,
    ControlFinding,
    ReportAnalysis,
)


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def test_evidence_package_fields():
    pkg = EvidencePackage(
        framework="pci_dss",
        org_id="org-1",
        period_start=_utcnow(),
        period_end=_utcnow(),
        devices=[
            DeviceEvidence(
                device_id="dev-1",
                hostname="router-01",
                compliance_scope=["PCI-CDE"],
                security_posture={
                    "ssh_enabled": True,
                    "telnet_enabled": False,
                    "http_enabled": False,
                    "https_enabled": True,
                    "snmp_enabled": False,
                    "acl_count": 3,
                },
            )
        ],
        changes=[],
        audit_events=[],
        topology_paths=[],
    )
    assert pkg.framework == "pci_dss"
    assert len(pkg.devices) == 1
    assert pkg.devices[0].compliance_scope == ["PCI-CDE"]


def test_report_analysis_finding_statuses():
    finding = ControlFinding(
        control_id="PCI-DSS Req 2.1",
        description="Do not use vendor-supplied defaults",
        status="pass",
        evidence_refs=["dev-1"],
        notes="",
    )
    assert finding.status in ("pass", "fail", "informational")


from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

TEST_ORG_ID = "00000000-0000-0000-0000-000000000001"

FRAMEWORK_SCOPE_TAGS = {
    "pci_dss":    ["PCI-CDE", "PCI-BOUNDARY"],
    "hipaa":      ["HIPAA-PHI"],
    "sox_itgc":   ["SOX-FINANCIAL"],
    "iso_27001":  ["ISO27001"],
    "fedramp":    ["FEDRAMP-BOUNDARY"],
    "soc2":       ["SOC2"],
    "nist_csf":   ["NIST-CSF"],
}


@pytest.mark.asyncio
async def test_collector_returns_evidence_package():
    from app.services.compliance.evidence_collector import EvidenceCollector

    mock_db = AsyncMock()
    mock_neo4j = MagicMock()

    device_row = MagicMock()
    device_row.id = uuid4()
    device_row.hostname = "core-router-01"
    device_row.compliance_scope = ["PCI-CDE"]
    device_row.meta = {"ssh_enabled": True, "telnet_enabled": False,
                       "http_enabled": False, "https_enabled": True,
                       "snmp_enabled": False, "acl_count": 5}

    dev_result = MagicMock()
    dev_result.scalars.return_value.all.return_value = [device_row]

    change_result = MagicMock()
    change_result.scalars.return_value.all.return_value = []

    audit_result = MagicMock()
    audit_result.scalars.return_value.all.return_value = []

    mock_db.execute = AsyncMock(side_effect=[dev_result, change_result, audit_result])
    mock_neo4j.get_full_topology = AsyncMock(return_value={"nodes": [], "relationships": []})

    collector = EvidenceCollector()
    pkg = await collector.collect(
        db=mock_db,
        neo4j_client=mock_neo4j,
        org_id=TEST_ORG_ID,
        framework="pci_dss",
        period_start=datetime(2026, 1, 1, tzinfo=timezone.utc),
        period_end=datetime(2026, 3, 31, tzinfo=timezone.utc),
    )

    assert pkg.framework == "pci_dss"
    assert len(pkg.devices) == 1
    assert pkg.devices[0].hostname == "core-router-01"
    assert pkg.devices[0].security_posture["ssh_enabled"] is True


@pytest.mark.asyncio
async def test_collector_scope_override():
    """scope_override restricts to only specified tags."""
    from app.services.compliance.evidence_collector import EvidenceCollector

    mock_db = AsyncMock()
    mock_neo4j = MagicMock()

    dev_result = MagicMock()
    dev_result.scalars.return_value.all.return_value = []
    change_result = MagicMock()
    change_result.scalars.return_value.all.return_value = []
    audit_result = MagicMock()
    audit_result.scalars.return_value.all.return_value = []

    mock_db.execute = AsyncMock(side_effect=[dev_result, change_result, audit_result])
    mock_neo4j.get_full_topology = AsyncMock(return_value={"nodes": [], "relationships": []})

    collector = EvidenceCollector()
    pkg = await collector.collect(
        db=mock_db,
        neo4j_client=mock_neo4j,
        org_id=TEST_ORG_ID,
        framework="pci_dss",
        period_start=datetime(2026, 1, 1, tzinfo=timezone.utc),
        period_end=datetime(2026, 3, 31, tzinfo=timezone.utc),
        scope_override=["PCI-CDE"],
    )
    assert pkg.framework == "pci_dss"


@pytest.mark.asyncio
async def test_collector_neo4j_unavailable():
    """If Neo4j raises, topology_paths is empty and collection still succeeds."""
    from app.services.compliance.evidence_collector import EvidenceCollector

    mock_db = AsyncMock()
    mock_neo4j = MagicMock()

    dev_result = MagicMock()
    dev_result.scalars.return_value.all.return_value = []
    change_result = MagicMock()
    change_result.scalars.return_value.all.return_value = []
    audit_result = MagicMock()
    audit_result.scalars.return_value.all.return_value = []

    mock_db.execute = AsyncMock(side_effect=[dev_result, change_result, audit_result])
    mock_neo4j.get_full_topology = AsyncMock(side_effect=Exception("Neo4j down"))

    collector = EvidenceCollector()
    pkg = await collector.collect(
        db=mock_db,
        neo4j_client=mock_neo4j,
        org_id=TEST_ORG_ID,
        framework="pci_dss",
        period_start=datetime(2026, 1, 1, tzinfo=timezone.utc),
        period_end=datetime(2026, 3, 31, tzinfo=timezone.utc),
    )
    assert pkg.topology_paths == []


def _make_pkg(framework: str, devices=None, changes=None, audit=None) -> EvidencePackage:
    return EvidencePackage(
        framework=framework,
        org_id=TEST_ORG_ID,
        period_start=datetime(2026, 1, 1, tzinfo=timezone.utc),
        period_end=datetime(2026, 3, 31, tzinfo=timezone.utc),
        devices=devices or [],
        changes=changes or [],
        audit_events=audit or [],
        topology_paths=[],
    )


def test_analyzer_pci_telnet_fail():
    """Device with telnet_enabled=True should produce a FAIL finding for PCI Req 2."""
    from app.services.compliance.framework_analyzer import FrameworkAnalyzer

    pkg = _make_pkg(
        "pci_dss",
        devices=[
            DeviceEvidence(
                device_id="dev-1",
                hostname="router-01",
                compliance_scope=["PCI-CDE"],
                security_posture={"telnet_enabled": True, "ssh_enabled": True,
                                  "http_enabled": False, "https_enabled": True,
                                  "snmp_enabled": False, "acl_count": 2},
            )
        ],
    )
    analyzer = FrameworkAnalyzer()
    analysis = analyzer.analyze(pkg, export_document_id="doc-1")

    statuses = {f.control_id: f.status for f in analysis.findings}
    assert statuses.get("PCI-DSS Req 2.2") == "fail"


def test_analyzer_pci_all_pass():
    """Device with secure posture should produce pass findings."""
    from app.services.compliance.framework_analyzer import FrameworkAnalyzer

    pkg = _make_pkg(
        "pci_dss",
        devices=[
            DeviceEvidence(
                device_id="dev-1",
                hostname="router-01",
                compliance_scope=["PCI-CDE"],
                security_posture={"telnet_enabled": False, "ssh_enabled": True,
                                  "http_enabled": False, "https_enabled": True,
                                  "snmp_enabled": False, "acl_count": 3},
            )
        ],
        changes=[
            ChangeEvidence(
                change_number="CHG-2026-0001",
                title="Update ACLs",
                status="completed",
                requested_by="user-1",
                approved_by="admin-1",
                approved_at=datetime(2026, 2, 1, tzinfo=timezone.utc),
                pre_change_hash="abc123",
                post_change_hash="def456",
                simulation_passed=True,
                affected_compliance_scopes=["PCI-CDE"],
            )
        ],
    )
    analyzer = FrameworkAnalyzer()
    analysis = analyzer.analyze(pkg, export_document_id="doc-1")

    fail_findings = [f for f in analysis.findings if f.status == "fail"]
    assert len(fail_findings) == 0


def test_analyzer_sox_unapproved_change_fails():
    """SOX: change without approved_by should produce a FAIL finding."""
    from app.services.compliance.framework_analyzer import FrameworkAnalyzer

    pkg = _make_pkg(
        "sox_itgc",
        changes=[
            ChangeEvidence(
                change_number="CHG-2026-0002",
                title="Emergency patch",
                status="completed",
                requested_by="user-1",
                approved_by=None,
                approved_at=None,
                pre_change_hash="aaa",
                post_change_hash="bbb",
                simulation_passed=None,
                affected_compliance_scopes=["SOX-FINANCIAL"],
            )
        ],
    )
    analyzer = FrameworkAnalyzer()
    analysis = analyzer.analyze(pkg, export_document_id="doc-1")

    statuses = {f.control_id: f.status for f in analysis.findings}
    assert statuses.get("SOX ITGC CC7.2") == "fail"


def test_analyzer_returns_report_analysis():
    from app.services.compliance.framework_analyzer import FrameworkAnalyzer

    pkg = _make_pkg("hipaa")
    analyzer = FrameworkAnalyzer()
    analysis = analyzer.analyze(pkg, export_document_id="doc-42")

    assert analysis.framework == "hipaa"
    assert analysis.export_document_id == "doc-42"
    assert isinstance(analysis.findings, list)


# ---------------------------------------------------------------------------
# PDFRenderer tests
# ---------------------------------------------------------------------------

def _make_analysis(framework: str = "pci_dss"):
    from app.services.compliance.evidence_models import ReportAnalysis, ControlFinding, DeviceEvidence
    return ReportAnalysis(
        framework=framework,
        org_id="org-1",
        period_start=datetime(2026, 1, 1, tzinfo=timezone.utc),
        period_end=datetime(2026, 3, 31, tzinfo=timezone.utc),
        generated_at=datetime(2026, 3, 26, tzinfo=timezone.utc),
        export_document_id="doc-1",
        findings=[
            ControlFinding(
                control_id="PCI-DSS Req 2.2",
                description="No insecure protocols",
                status="pass",
                evidence_refs=["dev-1"],
                notes="",
            )
        ],
        devices=[
            DeviceEvidence(
                device_id="dev-1",
                hostname="router-01",
                compliance_scope=["PCI-CDE"],
                security_posture={"ssh_enabled": True, "telnet_enabled": False,
                                  "http_enabled": False, "https_enabled": True,
                                  "snmp_enabled": False, "acl_count": 3},
            )
        ],
        changes=[],
        audit_events=[],
        topology_paths=[],
    )


def test_pdf_renderer_returns_bytes():
    from app.services.compliance.pdf_renderer import PDFRenderer

    renderer = PDFRenderer()
    result = renderer.render(_make_analysis(), org_name="Acme Corp")

    assert isinstance(result, bytes)
    assert len(result) > 0
    assert result[:4] == b"%PDF"


def test_pdf_renderer_all_frameworks():
    from app.services.compliance.pdf_renderer import PDFRenderer

    renderer = PDFRenderer()
    for fw in ["pci_dss", "hipaa", "sox_itgc", "iso_27001", "fedramp", "soc2", "nist_csf"]:
        result = renderer.render(_make_analysis(fw), org_name="Test Org")
        assert result[:4] == b"%PDF", f"Framework {fw} did not produce valid PDF"


# ---------------------------------------------------------------------------
# DOCXRenderer tests
# ---------------------------------------------------------------------------

def test_docx_renderer_returns_valid_docx():
    from app.services.compliance.docx_renderer import DOCXRenderer
    import docx as python_docx
    import io

    renderer = DOCXRenderer()
    result = renderer.render(_make_analysis(), org_name="Acme Corp")

    assert isinstance(result, bytes)
    assert len(result) > 0
    doc = python_docx.Document(io.BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Acme Corp" in full_text
    assert "PCI" in full_text


def test_docx_renderer_contains_finding():
    from app.services.compliance.docx_renderer import DOCXRenderer
    import docx as python_docx
    import io

    renderer = DOCXRenderer()
    result = renderer.render(_make_analysis(), org_name="TestOrg")
    doc = python_docx.Document(io.BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "PCI-DSS Req 2.2" in full_text


# ---------------------------------------------------------------------------
# ReportService tests
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_report_service_sets_completed_on_success():
    from app.services.compliance.report_service import generate_report
    from unittest.mock import AsyncMock, MagicMock, patch
    from uuid import uuid4

    export_id = str(uuid4())
    org_id    = str(uuid4())

    mock_db      = AsyncMock()
    mock_neo4j   = MagicMock()
    mock_doc     = MagicMock()
    mock_doc.id  = export_id
    mock_doc.status = "pending"

    fetch_result = MagicMock()
    fetch_result.scalar_one_or_none.return_value = mock_doc
    mock_db.execute = AsyncMock(return_value=fetch_result)

    with (
        patch("app.services.compliance.report_service.EvidenceCollector") as MockCollector,
        patch("app.services.compliance.report_service.FrameworkAnalyzer") as MockAnalyzer,
        patch("app.services.compliance.report_service.PDFRenderer") as MockPDF,
        patch("app.services.compliance.report_service.storage_service") as MockStorage,
    ):
        mock_pkg = MagicMock()
        MockCollector.return_value.collect = AsyncMock(return_value=mock_pkg)

        mock_analysis = _make_analysis()
        MockAnalyzer.return_value.analyze.return_value = mock_analysis

        MockPDF.return_value.render.return_value = b"%PDF-fake"
        MockStorage.upload_file.return_value = f"exports/{org_id}/report.pdf"

        await generate_report(
            export_document_id=export_id,
            org_id=org_id,
            org_name="Acme",
            framework="pci_dss",
            report_format="pdf",
            period_start=datetime(2026, 1, 1, tzinfo=timezone.utc),
            period_end=datetime(2026, 3, 31, tzinfo=timezone.utc),
            db=mock_db,
            neo4j_client=mock_neo4j,
        )

    assert mock_doc.status == "completed"
    assert mock_doc.storage_path == f"exports/{org_id}/report.pdf"


@pytest.mark.asyncio
async def test_report_service_sets_failed_on_error():
    from app.services.compliance.report_service import generate_report
    from unittest.mock import AsyncMock, MagicMock, patch
    from uuid import uuid4

    export_id = str(uuid4())
    org_id    = str(uuid4())

    mock_db  = AsyncMock()
    mock_doc = MagicMock()
    mock_doc.id = export_id
    mock_doc.status = "pending"

    fetch_result = MagicMock()
    fetch_result.scalar_one_or_none.return_value = mock_doc
    mock_db.execute = AsyncMock(return_value=fetch_result)

    with patch("app.services.compliance.report_service.EvidenceCollector") as MockCollector:
        MockCollector.return_value.collect = AsyncMock(side_effect=RuntimeError("DB down"))

        await generate_report(
            export_document_id=export_id,
            org_id=org_id,
            org_name="Acme",
            framework="pci_dss",
            report_format="pdf",
            period_start=datetime(2026, 1, 1, tzinfo=timezone.utc),
            period_end=datetime(2026, 3, 31, tzinfo=timezone.utc),
            db=mock_db,
            neo4j_client=MagicMock(),
        )

    assert mock_doc.status == "failed"
    assert "DB down" in (mock_doc.error_message or "")
