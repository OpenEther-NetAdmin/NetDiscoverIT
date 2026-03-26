"""
DOCXRenderer — renders a ReportAnalysis to DOCX bytes using python-docx.
"""
from __future__ import annotations

import io

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.compliance.evidence_models import ControlFinding, ReportAnalysis

_FRAMEWORK_LABELS = {
    "pci_dss":   "PCI-DSS v4.0",
    "hipaa":     "HIPAA Security Rule",
    "sox_itgc":  "SOX ITGC",
    "iso_27001": "ISO 27001:2022",
    "fedramp":   "FedRAMP Moderate",
    "soc2":      "SOC 2 Type II",
    "nist_csf":  "NIST Cybersecurity Framework",
}

_STATUS_COLORS = {
    "pass":          RGBColor(0x2e, 0x7d, 0x32),
    "fail":          RGBColor(0xc6, 0x28, 0x28),
    "informational": RGBColor(0x15, 0x65, 0xc0),
}


class DOCXRenderer:
    """Renders a ReportAnalysis to DOCX bytes."""

    def render(self, analysis: ReportAnalysis, org_name: str = "Organization") -> bytes:
        doc = Document()

        self._cover(doc, analysis, org_name)
        self._executive_summary(doc, analysis)
        self._scope_section(doc, analysis)
        self._findings_section(doc, analysis)
        self._changes_section(doc, analysis)
        self._audit_section(doc, analysis)
        self._appendix(doc, analysis)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _cover(self, doc: Document, analysis: ReportAnalysis, org_name: str) -> None:
        label = _FRAMEWORK_LABELS.get(analysis.framework, analysis.framework.upper())
        p = doc.add_heading(org_name, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading(f"{label} Compliance Report", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        period = (
            f"{analysis.period_start.strftime('%Y-%m-%d')} — "
            f"{analysis.period_end.strftime('%Y-%m-%d')}"
        )
        doc.add_paragraph(f"Assessment Period: {period}")
        doc.add_paragraph(f"Generated: {analysis.generated_at.strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph("CONFIDENTIAL — RESTRICTED DISTRIBUTION")
        doc.add_page_break()

    def _executive_summary(self, doc: Document, analysis: ReportAnalysis) -> None:
        doc.add_heading("Executive Summary", 1)
        passed = sum(1 for f in analysis.findings if f.status == "pass")
        failed = sum(1 for f in analysis.findings if f.status == "fail")
        info   = sum(1 for f in analysis.findings if f.status == "informational")
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Value"
        for label, val in [
            ("Controls Assessed", len(analysis.findings)),
            ("PASS", passed),
            ("FAIL", failed),
            ("INFORMATIONAL", info),
            ("In-Scope Devices", len(analysis.devices)),
            ("Changes Reviewed", len(analysis.changes)),
        ]:
            row = tbl.add_row().cells
            row[0].text = label
            row[1].text = str(val)

    def _scope_section(self, doc: Document, analysis: ReportAnalysis) -> None:
        doc.add_heading("Scope — In-Scope Devices", 2)
        if not analysis.devices:
            doc.add_paragraph("No in-scope devices found for this framework.")
            return
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Hostname"
        hdr[1].text = "Device ID"
        hdr[2].text = "Compliance Scope"
        for d in analysis.devices:
            row = tbl.add_row().cells
            row[0].text = d.hostname or "(unknown)"
            row[1].text = d.device_id[:8] + "..."
            row[2].text = ", ".join(d.compliance_scope)

    def _findings_section(self, doc: Document, analysis: ReportAnalysis) -> None:
        doc.add_heading("Control Findings", 2)
        for finding in analysis.findings:
            color = _STATUS_COLORS.get(finding.status, RGBColor(0, 0, 0))
            p = doc.add_paragraph()
            run = p.add_run(f"[{finding.status.upper()}] {finding.control_id}")
            run.bold = True
            run.font.color.rgb = color
            doc.add_paragraph(finding.description)
            if finding.notes:
                doc.add_paragraph(finding.notes).italic = True  # type: ignore[assignment]
            if finding.evidence_refs:
                refs = ", ".join(finding.evidence_refs[:5])
                doc.add_paragraph(f"Evidence: {refs}")

    def _changes_section(self, doc: Document, analysis: ReportAnalysis) -> None:
        if not analysis.changes:
            return
        doc.add_heading("Change Management Evidence", 2)
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        for i, h in enumerate(["CHG #", "Title", "Approved By", "Approved At", "Sim", "Status"]):
            tbl.rows[0].cells[i].text = h
        for c in analysis.changes:
            row = tbl.add_row().cells
            row[0].text = c.change_number
            row[1].text = (c.title[:40] + "...") if len(c.title) > 40 else c.title
            row[2].text = c.approved_by[:8] + "..." if c.approved_by else "—"
            row[3].text = c.approved_at.strftime("%Y-%m-%d") if c.approved_at else "—"
            row[4].text = "✓" if c.simulation_passed else ("—" if c.simulation_passed is None else "✗")
            row[5].text = c.status

    def _audit_section(self, doc: Document, analysis: ReportAnalysis) -> None:
        if not analysis.audit_events:
            return
        doc.add_heading("Audit Trail Summary", 2)
        doc.add_paragraph(f"{len(analysis.audit_events)} audit events in assessment period.")
        from collections import Counter
        counts = Counter(e.action for e in analysis.audit_events)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.rows[0].cells[0].text = "Action"
        tbl.rows[0].cells[1].text = "Count"
        for action, count in counts.most_common(10):
            row = tbl.add_row().cells
            row[0].text = action
            row[1].text = str(count)

    def _appendix(self, doc: Document, analysis: ReportAnalysis) -> None:
        doc.add_heading("Appendix — Report Metadata", 2)
        doc.add_paragraph(f"Export Document ID: {analysis.export_document_id}")
        doc.add_paragraph(f"Organization ID: {analysis.org_id}")
        doc.add_paragraph(f"Framework: {analysis.framework}")
        doc.add_paragraph(f"Generated at: {analysis.generated_at.isoformat()}")
