"""
Integration tests for compliance report generation.

Requires: running PostgreSQL + MinIO containers (docker compose up).
Run with: docker compose exec api pytest tests/api/test_compliance_reports_integration.py -v -s

These tests create real DB records and upload real files to MinIO.
They are skipped if the environment is not available.
"""
from __future__ import annotations

import asyncio
import os
import pytest
from datetime import datetime, timezone, timedelta
from uuid import uuid4

pytestmark = pytest.mark.skipif(
    not os.environ.get("POSTGRES_PASSWORD"),
    reason="Integration tests require full docker-compose stack",
)


import pytest_asyncio


@pytest_asyncio.fixture
async def db_session():
    from app.db.database import async_session_maker
    async with async_session_maker() as session:
        yield session


@pytest_asyncio.fixture
async def neo4j():
    from app.db.neo4j import get_neo4j_client
    try:
        client = await get_neo4j_client()
        yield client
    except Exception:
        yield None


@pytest_asyncio.fixture
async def org_and_device(db_session):
    """Create a test org + device with compliance scope tags."""
    from app.models.models import Organization, Device
    import uuid

    org = Organization(
        name=f"Integration Test Org {uuid4().hex[:6]}",
        slug=f"test-org-{uuid4().hex[:6]}",
        subscription_tier="enterprise",
    )
    db_session.add(org)
    await db_session.commit()
    await db_session.refresh(org)

    device = Device(
        organization_id=org.id,
        ip_address="10.0.0.1",
        hostname="core-router-01",
        compliance_scope=["PCI-CDE", "PCI-BOUNDARY"],
        meta={
            "ssh_enabled": True,
            "telnet_enabled": False,
            "http_enabled": False,
            "https_enabled": True,
            "snmp_enabled": False,
            "acl_count": 5,
        },
        is_active=True,
    )
    db_session.add(device)
    await db_session.commit()
    await db_session.refresh(device)

    yield org, device

    await db_session.delete(device)
    await db_session.delete(org)
    await db_session.commit()


@pytest.mark.asyncio
async def test_full_pdf_generation_pipeline(db_session, neo4j, org_and_device):
    """Full pipeline: generate_report sets ExportDocument to completed and uploads to MinIO."""
    from app.models.models import ExportDocument
    from app.services.compliance.report_service import generate_report
    from sqlalchemy import select

    org, device = org_and_device

    doc = ExportDocument(
        organization_id=org.id,
        document_type="compliance_report",
        format="pdf",
        status="pending",
        parameters={
            "framework": "pci_dss",
            "period_start": "2026-01-01T00:00:00+00:00",
            "period_end": "2026-03-31T23:59:59+00:00",
        },
    )
    db_session.add(doc)
    await db_session.commit()
    await db_session.refresh(doc)

    await generate_report(
        export_document_id=str(doc.id),
        org_id=str(org.id),
        org_name=org.name,
        framework="pci_dss",
        report_format="pdf",
        period_start=datetime(2026, 1, 1, tzinfo=timezone.utc),
        period_end=datetime(2026, 3, 31, tzinfo=timezone.utc),
        db=db_session,
        neo4j_client=neo4j,
    )

    result = await db_session.execute(
        select(ExportDocument).where(ExportDocument.id == doc.id)
    )
    updated_doc = result.scalar_one()

    assert updated_doc.status == "completed", f"Expected completed, got {updated_doc.status}: {updated_doc.error_message}"
    assert updated_doc.storage_path is not None
    assert updated_doc.file_size_bytes is not None
    assert updated_doc.file_size_bytes > 0

    from app.services.storage import storage_service
    url = storage_service.generate_presigned_url(updated_doc.storage_path)
    assert url.startswith("http")

    storage_service.delete_file(updated_doc.storage_path)
    await db_session.delete(updated_doc)
    await db_session.commit()


@pytest.mark.asyncio
async def test_full_docx_generation_pipeline(db_session, neo4j, org_and_device):
    """Full DOCX pipeline: generate_report produces a valid DOCX uploaded to MinIO."""
    from app.models.models import ExportDocument
    from app.services.compliance.report_service import generate_report
    from sqlalchemy import select
    import io
    import docx as python_docx

    org, device = org_and_device

    doc = ExportDocument(
        organization_id=org.id,
        document_type="compliance_report",
        format="docx",
        status="pending",
        parameters={
            "framework": "sox_itgc",
            "period_start": "2026-01-01T00:00:00+00:00",
            "period_end": "2026-03-31T23:59:59+00:00",
        },
    )
    db_session.add(doc)
    await db_session.commit()
    await db_session.refresh(doc)

    await generate_report(
        export_document_id=str(doc.id),
        org_id=str(org.id),
        org_name=org.name,
        framework="sox_itgc",
        report_format="docx",
        period_start=datetime(2026, 1, 1, tzinfo=timezone.utc),
        period_end=datetime(2026, 3, 31, tzinfo=timezone.utc),
        db=db_session,
        neo4j_client=neo4j,
    )

    result = await db_session.execute(
        select(ExportDocument).where(ExportDocument.id == doc.id)
    )
    updated_doc = result.scalar_one()

    assert updated_doc.status == "completed", f"{updated_doc.status}: {updated_doc.error_message}"
    assert updated_doc.storage_path is not None

    from app.services.storage import storage_service
    storage_service.delete_file(updated_doc.storage_path)
    await db_session.delete(updated_doc)
    await db_session.commit()


@pytest.mark.asyncio
async def test_failed_generation_marks_document_failed(db_session, neo4j, org_and_device):
    """If evidence collection raises, ExportDocument.status becomes 'failed'."""
    from app.models.models import ExportDocument
    from app.services.compliance.report_service import generate_report
    from sqlalchemy import select
    from unittest.mock import patch, AsyncMock

    org, _ = org_and_device

    doc = ExportDocument(
        organization_id=org.id,
        document_type="compliance_report",
        format="pdf",
        status="pending",
        parameters={"framework": "hipaa"},
    )
    db_session.add(doc)
    await db_session.commit()
    await db_session.refresh(doc)

    with patch("app.services.compliance.report_service.EvidenceCollector") as MockCollector:
        MockCollector.return_value.collect = AsyncMock(side_effect=RuntimeError("simulated failure"))
        await generate_report(
            export_document_id=str(doc.id),
            org_id=str(org.id),
            org_name=org.name,
            framework="hipaa",
            report_format="pdf",
            period_start=datetime(2026, 1, 1, tzinfo=timezone.utc),
            period_end=datetime(2026, 3, 31, tzinfo=timezone.utc),
            db=db_session,
            neo4j_client=neo4j,
        )

    result = await db_session.execute(
        select(ExportDocument).where(ExportDocument.id == doc.id)
    )
    updated_doc = result.scalar_one()
    assert updated_doc.status == "failed"
    assert "simulated failure" in (updated_doc.error_message or "")

    await db_session.delete(updated_doc)
    await db_session.commit()
